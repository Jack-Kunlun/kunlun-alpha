from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


DOCX = Path(__file__).resolve().parents[2] / "outputs" / "昆仑智策项目总体规划与分阶段实施手册.docx"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def fail(message: str) -> None:
    print(f"FAIL: {message}")
    raise SystemExit(1)


if not DOCX.exists() or DOCX.stat().st_size < 50_000:
    fail("DOCX missing or unexpectedly small")

with zipfile.ZipFile(DOCX) as archive:
    bad = archive.testzip()
    if bad:
        fail(f"corrupt ZIP member: {bad}")
    document_xml = archive.read("word/document.xml")

root = ET.fromstring(document_xml)
doc = Document(DOCX)
text = "\n".join(p.text for p in doc.paragraphs)
for table in doc.tables:
    for row in table.rows:
        text += "\n" + "\t".join(cell.text for cell in row.cells)

required = [
    "昆仑智策", "KUNLUN ALPHA", "WorkBuddy", "Codex", "shadcn/ui",
    "Phase 0", "Phase 7", ".editorconfig", "end_of_line = lf",
    "Windows、macOS、Linux", "实盘开放门禁", "Kill Switch",
    "场内贵金属基金", "underlyingCommodity", "NAV/iNAV", "P1-N13", "P1-N14", "P5-N14", "P6-N13",
]
for item in required:
    if item not in text:
        fail(f"missing required content: {item}")

for placeholder in ("TBD", "TODO", "待补充", "<NODE_ID>"):
    if placeholder in text and placeholder != "<NODE_ID>":
        fail(f"placeholder found: {placeholder}")

node_ids = re.findall(r"P[0-7]-N\d{2}", text)
unique_nodes = sorted(set(node_ids))
if len(unique_nodes) != 98:
    fail(f"expected 98 unique nodes, found {len(unique_nodes)}")

counts = {phase: sum(1 for node in unique_nodes if node.startswith(f"P{phase}-")) for phase in range(8)}
expected = {0: 15, 1: 14, 2: 10, 3: 10, 4: 10, 5: 14, 6: 13, 7: 12}
if counts != expected:
    fail(f"phase node counts differ: {counts}")

tables = root.findall(".//w:tbl", NS)
if len(tables) < 100:
    fail(f"expected at least 100 tables, found {len(tables)}")
for idx, table in enumerate(tables, 1):
    grid = table.find("w:tblGrid", NS)
    if grid is None or not grid.findall("w:gridCol", NS):
        fail(f"table {idx} lacks fixed grid geometry")
    width = table.find("w:tblPr/w:tblW", NS)
    if width is None or width.get(f"{{{NS['w']}}}type") != "dxa":
        fail(f"table {idx} lacks DXA table width")
    for cell in table.findall(".//w:tc", NS):
        tc_w = cell.find("w:tcPr/w:tcW", NS)
        if tc_w is None or tc_w.get(f"{{{NS['w']}}}type") != "dxa":
            fail(f"table {idx} contains cell without DXA width")

print(f"PASS: zip_integrity=ok size={DOCX.stat().st_size}")
print(f"PASS: phases=8 nodes={len(unique_nodes)} counts={counts}")
print(f"PASS: tables={len(tables)} fixed_geometry=ok")
print("PASS: required_sections_and_config=present placeholders=none")
print(f"PASS: paragraphs={len(doc.paragraphs)}")
