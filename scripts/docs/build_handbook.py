from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "outputs" / "昆仑智策项目总体规划与分阶段实施手册.docx"

NAVY = "17324D"
BLUE = "2E668D"
TEAL = "247B7B"
GOLD = "B88746"
INK = "1F2933"
MUTED = "667786"
LIGHT = "EEF3F6"
PALE = "F7F9FA"
WHITE = "FFFFFF"
RED = "A23B3B"
GREEN = "2F6F55"


@dataclass(frozen=True)
class Node:
    node_id: str
    title: str
    objective: str
    files: str
    deliverable: str
    acceptance: str
    risk: str


@dataclass(frozen=True)
class Phase:
    number: int
    name: str
    goal: str
    exit_gate: str
    nodes: tuple[Node, ...]


def N(i: str, title: str, objective: str, files: str, deliverable: str, acceptance: str, risk: str) -> Node:
    return Node(i, title, objective, files, deliverable, acceptance, risk)


PHASES: tuple[Phase, ...] = (
    Phase(0, "工程基础设施", "建立可在 Windows、macOS、Linux 上一致运行的 Monorepo、质量门禁、基础服务与可观测性。", "新开发者能按文档完成启动；Web、API、Python 服务与基础设施健康；CI 全绿；仓库不存在任何实盘执行能力。", (
        N("P0-N01", "初始化仓库治理文件", "建立根目录、许可证、README、忽略规则与贡献约定。", "README.md、CONTRIBUTING.md、.gitignore、docs/adr/", "项目入口、开发原则、ADR 模板和目录责任说明。", "三端均可读取；无密钥或机器路径；文档链接有效。", "避免把暂定架构写成不可修改的永久约束。"),
        N("P0-N02", "统一 LF 与 EditorConfig", "用统一配置消除跨平台换行和缩进差异。", ".editorconfig、.gitattributes、scripts/check-line-endings.*", "LF/UTF-8 基线、Makefile 与批处理例外、CI 检查入口。", "故意写入 CRLF 的普通文本会被检查拒绝；三端行为一致。", "注意 Markdown 尾随空格和二进制文件误判。"),
        N("P0-N03", "锁定运行时与包管理器", "解析并锁定实施时最新稳定的 Node.js、pnpm、Python、uv。", ".node-version、package.json、pyproject.toml、.python-version、docs/versions.md", "版本基线、升级策略和验证命令。", "错误版本启动时得到清晰提示；禁止预发布和浮动版本。", "不要在后续业务节点顺带升级。"),
        N("P0-N04", "初始化 Turborepo 与 pnpm workspace", "建立 TypeScript 工作区和统一任务图。", "package.json、pnpm-workspace.yaml、turbo.json、packages/config/", "dev/build/test/lint/typecheck 任务及缓存边界。", "空工作区全量命令成功，缓存不包含密钥和运行态数据。", "任务依赖方向必须与模块边界一致。"),
        N("P0-N05", "初始化 uv Python workspace", "建立服务与领域包分离的 Python 工作区。", "pyproject.toml、uv.lock、python/packages/、services/", "统一 lint/typecheck/test 命令和最小示例包。", "uv sync 与测试在三端路径规则下可运行。", "禁止服务包反向成为通用领域依赖。"),
        N("P0-N06", "创建 React Web 骨架", "建立 Vite、React、严格 TypeScript 与路由外壳。", "apps/web/、packages/shared-types/", "可访问的应用壳、错误边界、路由和环境配置。", "加载、404、运行错误和无配置状态有测试。", "浏览器环境变量不得泄露服务端密钥。"),
        N("P0-N07", "建立 shadcn/ui 设计基线", "统一颜色、字体、间距、暗色模式和基础组件策略。", "apps/web/components/ui/、apps/web/styles/、components.json", "Button、Dialog、Form、Table、Toast 等受控基础组件。", "键盘、焦点、暗色和响应式检查通过。", "禁止建立与 shadcn/ui 竞争的第二套基础组件。"),
        N("P0-N08", "创建 NestJS API 骨架", "建立配置、健康检查、日志、异常和版本化路由。", "apps/api/src/、packages/contracts/", "健康端点、统一错误响应、请求 ID 和 OpenAPI 基线。", "启动、错误映射、配置缺失和关闭流程测试通过。", "控制面不能承担量化计算。"),
        N("P0-N09", "创建 Python 服务骨架", "为 Engine 和 Worker 提供一致生命周期。", "services/*/src/、python/packages/ashare-common/", "配置、日志、健康检查、优雅关闭和任务入口。", "服务可独立启动并在终止信号后安全退出。", "不要复制配置和日志实现。"),
        N("P0-N10", "建立跨语言契约流水线", "定义 JSON Schema/OpenAPI 并生成 TypeScript/Python 类型。", "packages/contracts/、python/packages/ashare-contracts/", "契约版本、生成命令、兼容性检查和示例载荷。", "生成结果可复现；破坏性变更会阻断 CI。", "手写类型不得与生成契约并行漂移。"),
        N("P0-N11", "搭建本地数据基础设施", "用 Docker Compose 启动 PostgreSQL、Redis、ClickHouse、MinIO。", "infra/docker/、infra/*/", "健康检查、持久卷、初始化脚本和示例配置。", "全新环境可一键启动、停止并重复初始化。", "默认凭证只能用于本地且必须显式标记。"),
        N("P0-N12", "搭建可观测性基础设施", "统一追踪、指标和日志关联。", "infra/prometheus/、infra/grafana/、infra/loki/、packages/observability/", "OpenTelemetry 接入、基础仪表盘和告警规则。", "一次请求可通过 trace/request ID 贯穿 Web、API 和 Python 服务。", "控制标签基数，避免证券代码造成指标爆炸。"),
        N("P0-N13", "配置统一质量门禁", "配置格式化、静态分析、单元测试和覆盖率。", "eslint、prettier、ruff、pyright、vitest、jest、pytest 配置", "根命令和包级命令，含失败示例验证。", "格式、类型、测试任一失败均阻断合并。", "避免仅为覆盖率数字编写无断言测试。"),
        N("P0-N14", "建立 CI 三端矩阵", "验证核心工具在 Linux、Windows、macOS 的一致性。", ".github/workflows/ci.yml、scripts/ci/", "缓存、分层任务、行尾检查和构建产物检查。", "Linux 全量通过；Windows/macOS 至少验证 bootstrap、lint、typecheck、核心测试。", "不要依赖 Bash 专属语法。"),
        N("P0-N15", "完成开发者启动手册", "让无上下文开发者可重复搭建环境。", "docs/development/、.env.example", "安装、启动、故障排查、端口与数据重置说明。", "在干净机器或容器中按文档完成演练。", "文档不得包含真实凭证或个人目录。"),
    )),
    Phase(1, "A 股与场内基金数据底座", "建立统一证券/基金主数据、交易日历、Provider 抽象、可追溯数据流水线与质量体系。", "数据可重放、校验和追溯；原始区不可变；股票、ETF、场内贵金属基金、时间、净值与复权口径统一；质量告警可见。", (
        N("P1-N01", "定义证券主数据契约", "定义 Instrument、Exchange、Board、TradingStatus 与统一代码。", "packages/contracts/instrument/、python/packages/market-core/", "跨语言模型、代码解析器和示例数据。", "有效与无效代码、退市、ST、ETF 和板块规则测试通过。", "禁止多种代码格式进入领域层。"),
        N("P1-N02", "定义交易日历契约", "统一交易日、交易时段、节假日和临时休市。", "contracts/calendar/、market-core/calendar/", "CalendarProvider、会话查询和时区规则。", "跨午夜、午休、节假日和边界时刻测试通过。", "不可用自然日替代交易日。"),
        N("P1-N03", "定义行情与公司行动契约", "定义 OHLCV、逐笔/分钟边界、复权因子和公司行动。", "contracts/market-data/、market-core/models/", "精度明确的行情模型与校验规则。", "价格、数量、金额、停牌和复权样例通过。", "研究价格与真实成交价格必须分离。"),
        N("P1-N04", "建立 Provider 接口族", "为证券、日历、行情、板块、新闻等定义独立适配器。", "market-core/providers/、ashare-contracts/providers/", "能力声明、错误分类、限流和游标接口。", "模拟 Provider 的契约测试覆盖超时、限流和空页。", "引擎不得依赖具体数据商 SDK。"),
        N("P1-N05", "实现原始数据落地区", "把每次外部响应按来源和日期不可变保存。", "services/data-worker/raw/、infra/minio/", "RawObjectManifest、校验和、压缩和重放入口。", "重复抓取不会覆盖原始对象；清单可定位来源与请求。", "原始区不得混入规范化结果。"),
        N("P1-N06", "实现证券主数据采集", "通过首个 Provider 完成全量与增量证券同步。", "services/data-worker/jobs/instruments/", "采集任务、检查点、失败重试和差异报告。", "重复运行幂等；新增、变更、退市均有审计记录。", "不得以名称作为证券唯一键。"),
        N("P1-N07", "实现交易日历采集", "同步交易所日历并提供缺口检测。", "services/data-worker/jobs/calendars/", "日历版本、来源追踪和手工修正审计。", "缺失日期、冲突来源和临时休市产生告警。", "手工修正不能静默覆盖来源数据。"),
        N("P1-N08", "实现行情规范化与校验", "统一字段、精度、时区和异常分类。", "data-worker/normalize/、market-core/validation/", "Normalizer、Validator、拒绝区和质量事件。", "乱序、重复、负价格、异常成交量与缺列测试通过。", "异常数据不得悄悄丢弃。"),
        N("P1-N09", "建立 PostgreSQL 业务模型", "保存 Provider、任务、检查点、数据版本和质量事件。", "apps/api/migrations/、packages/db/", "迁移、回滚和仓储接口。", "空库升级、回滚、重复迁移和约束测试通过。", "不要把高频行情写入 PostgreSQL。"),
        N("P1-N10", "建立 ClickHouse 行情模型", "设计分钟/日线、公司行动和查询分区。", "infra/clickhouse/migrations/、market-core/storage/", "表结构、排序键、TTL 策略和写入器。", "批量写入幂等；典型区间查询达到基线。", "排序键需匹配主要查询，避免过度分区。"),
        N("P1-N11", "实现增量调度与恢复", "调度采集任务并从检查点恢复。", "data-worker/scheduler/、apps/api/jobs/", "任务状态机、租约、退避和死信处理。", "进程中断后不重写已确认数据，失败可重放。", "重试必须区分永久错误与暂时错误。"),
        N("P1-N12", "建立数据质量与对账面板", "可视化覆盖率、延迟、重复、缺口和来源差异。", "apps/api/modules/data-quality/、apps/web/features/data-quality/", "质量 API、告警和 shadcn/ui 状态页。", "可按日期、来源、证券定位问题并链接原始证据。", "面板不能只展示总量而隐藏异常明细。"),
        N("P1-N13", "定义场内贵金属基金契约", "在统一 Instrument/基金模型中定义黄金及白银相关场内 ETF/基金分类。", "packages/contracts/funds/、python/packages/market-core/funds/", "fundAssetClass、underlyingCommodity、币种、基准、费率、有效期和来源契约。", "GOLD/SILVER/OTHER、未知分类、币种不一致和历史有效期测试通过。", "不得仅凭产品名称猜测底层商品，也不得引入现货或期货语义。"),
        N("P1-N14", "采集贵金属基金元数据与净值", "采集、校验、存储并监控场内贵金属基金元数据、NAV 和 iNAV。", "services/data-worker/jobs/precious-metals-funds/、market-core/funds/validation/", "Provider 能力、原始清单、规范化记录、溢折价校验、时效告警和质量面板。", "负值、过期 NAV/iNAV、来源冲突、重复记录和可用时间边界测试通过。", "NAV/iNAV 不能被标记为可成交价格，缺失字段不得静默推断。"),
    )),
    Phase(2, "市场情绪与板块轮动", "把 A 股涨跌停、连板、溢价和板块资金变化转化为确定性、版本化、可回放的市场状态。", "同一输入可重算相同结果；算法版本可并存；分钟快照与轮动事件能历史回放。", (
        N("P2-N01", "定义涨跌停事实模型", "建模涨停、跌停、炸板、首封、末封和开板次数。", "emotion-core/models/、contracts/emotion/", "LimitEvent 与 LimitPoolSnapshot 契约。", "ST、不同板块涨跌幅和无价格情形测试通过。", "涨停判断必须基于当日规则而非固定百分比。"),
        N("P2-N02", "构建涨跌停池计算器", "从规范化行情生成分钟事实与日终事实。", "emotion-core/limit_pool/", "纯函数计算器、增量聚合器和测试夹具。", "重放与实时输入结果一致；乱序事件可纠正。", "不可把展示排序混入事实计算。"),
        N("P2-N03", "实现连板梯队与晋级率", "计算几天几板、最高板、1进2 等指标。", "emotion-core/ladder/", "BoardLadderSnapshot 与历史晋级统计。", "停牌、断板、跨交易日和新股样例正确。", "交易日连续性不能用自然日判断。"),
        N("P2-N04", "实现昨日溢价与亏钱效应", "计算昨日涨停/连板表现和高位回撤。", "emotion-core/premium/", "可配置观察窗口和无偏样本规则。", "幸存者偏差、停牌和一字板处理有测试。", "不得使用当时不可获得的收盘后字段。"),
        N("P2-N05", "定义情绪评分 V1", "用确定性标准化和权重组合生成分项与总分。", "emotion-core/scoring/v1/、docs/algorithms/", "emotion_score_v1、权重配置和解释字段。", "固定样本得到固定分数；缺失数据降级规则明确。", "版本升级不得覆盖旧分数。"),
        N("P2-N06", "建立板块主数据与映射", "统一行业、概念、题材、股票与 ETF 关系。", "rotation-core/taxonomy/、contracts/sector/", "层级、别名、有效期和来源优先级。", "多归属、别名冲突和历史变更可追溯。", "不要把临时热点永久写入固定行业分类。"),
        N("P2-N07", "计算 SectorSnapshot", "按分钟计算涨幅、速度、成交额、宽度、龙头与强度。", "rotation-core/snapshot/", "版本化 SectorSnapshot 与批流一致接口。", "成分变更、缺行情和异常成交量测试通过。", "聚合必须使用当时有效的成分集合。"),
        N("P2-N08", "定义板块生命周期状态", "输出启动、发酵、加速、高潮、分歧、修复、退潮、冰点。", "rotation-core/lifecycle/v1/", "状态机、阈值配置、转换原因和置信度。", "非法跳转被拒绝；回放结果可解释。", "状态标签不能只由单一涨幅阈值决定。"),
        N("P2-N09", "检测 RotationEvent", "识别板块强度首次交叉、持续领先和快速衰退。", "rotation-core/events/", "from/to、时间、置信度、证据窗口。", "抖动去噪、并列、短暂超越和回撤测试通过。", "避免每分钟重复产生同一事件。"),
        N("P2-N10", "提供情绪与轮动回放终端", "通过 API 和 shadcn/ui 展示快照、梯队和时间轴。", "apps/api/modules/emotion/、apps/web/features/market-regime/", "查询 API、分钟时间轴、筛选与错误状态。", "任意历史交易日可回放并展示算法版本。", "实时与历史接口的口径必须一致。"),
    )),
    Phase(3, "新闻、事件与热点 Intelligence", "将新闻、政策、公告和社交内容转换为带证据、版本和置信度的事件与热点。", "任一事件可追溯原文和模型版本；离线评测可重复；AI 失败不会污染事实层。", (
        N("P3-N01", "定义 RawContent 与来源策略", "统一新闻、公告、研报、互动和社交内容。", "event-core/content/、contracts/content/", "来源、发布时间、抓取时间、正文指纹和许可元数据。", "重复 URL、转载、更新时间和删除状态测试通过。", "必须保留来源和使用限制。"),
        N("P3-N02", "实现内容规范化", "清理编码、HTML、时间、作者和附件元数据。", "intelligence-engine/normalize/", "可逆清洗记录和规范化文本。", "乱码、空正文、附件和超长内容有确定处理。", "不能丢失影响证据定位的原始偏移信息。"),
        N("P3-N03", "实现近重复检测", "组合哈希、标题和语义相似度识别转载。", "event-core/dedup/", "内容簇、主记录和阈值评估集。", "相同事件不同来源可聚类且保留各自证据。", "不要把不同公司同类公告误合并。"),
        N("P3-N04", "建立实体识别与证券解析", "抽取公司、证券、机构、人物、政策和地区。", "event-core/entities/、market-core/resolver/", "候选、消歧、置信度和人工覆盖。", "简称、曾用名、同名公司和多证券映射测试通过。", "低置信度不得强制绑定股票。"),
        N("P3-N05", "定义 MarketEvent 契约", "统一事件类型、情绪、影响、主题、证券和证据。", "contracts/event/、event-core/models/", "版本化事件模式与兼容性样例。", "每个结论均能链接证据片段和来源。", "模型输出不能绕过结构校验。"),
        N("P3-N06", "建立 Prompt 与模型注册表", "管理提示词、模型、参数、成本和输出模式。", "intelligence-engine/llm/registry/、docs/prompts/", "不可变版本、灰度配置和回滚能力。", "同一版本可复现实验；敏感配置不入库。", "不得用未登记提示词直接写生产数据。"),
        N("P3-N07", "实现事件抽取流水线", "批量/实时抽取类型、情绪、影响和主题映射。", "intelligence-engine/pipelines/events/", "队列消费者、验证、重试、拒绝区和审计。", "无效 JSON、超时、限流和低置信度路径测试通过。", "AI 错误必须可隔离和重跑。"),
        N("P3-N08", "实现 HotTopic 聚合", "按窗口聚合事件、市场表现、成交额和社交热度。", "hotspot-core/aggregation/", "HotTopic、HotScore V1 和生命周期状态。", "窗口边界、重复内容和稀疏数据结果稳定。", "事件数量不能直接等同热点强度。"),
        N("P3-N09", "建立离线评测与反馈", "用标注集评估实体、事件、情绪和映射质量。", "research/evaluations/events/、apps/api/modules/feedback/", "指标、错误切片、反馈闭环和版本对比。", "评测数据冻结；新版本必须与基线对比。", "避免把线上反馈直接作为未经审核的真值。"),
        N("P3-N10", "构建事件与热点工作台", "展示热点排行、事件时间轴、证据和置信度。", "apps/web/features/intelligence/、apps/api/modules/intelligence/", "shadcn/ui 列表、详情、筛选、反馈和证据链接。", "无证据结果不可显示为确定事实；四类页面状态完整。", "UI 必须区分事实、推断和低置信度结论。"),
    )),
    Phase(4, "龙虎榜席位与 KOL Intelligence", "建立可审计的席位知识库和 KOL 公开 Claim 体系，并产生可回测特征。", "别名解析、证据等级与绩效计算均可追溯；Claim 与真实成交语义严格分离。", (
        N("P4-N01", "定义席位领域模型", "建模 SeatProfile、Alias、Trade、Performance 和 Preference。", "seat-core/models/、contracts/seat/", "版本化契约、稳定 seat_id 和有效期。", "名称变更、机构席位、未知席位和合并拆分测试通过。", "原始营业部名称不能作为唯一 ID。"),
        N("P4-N02", "实现席位别名解析", "用规则、相似度和人工审核统一多来源名称。", "seat-core/resolution/", "候选匹配、置信度、审核队列和审计日志。", "黄金样本准确率达到文档阈值，低置信度不自动合并。", "错误合并会污染全部历史统计。"),
        N("P4-N03", "实现龙虎榜采集与规范化", "采集席位买卖、上榜原因和证券事实。", "data-worker/jobs/seats/、seat-core/normalize/", "不可变原始记录、规范化成交和来源证据。", "重复榜单、修订、净额与买卖分项对账通过。", "净买入不能替代原始买卖明细。"),
        N("P4-N04", "计算席位绩效特征", "计算 T+1/T+3/T+5、胜率、回撤和风格偏好。", "seat-core/features/", "带样本数、窗口和口径版本的特征。", "停牌、涨跌停、复权和小样本处理有测试。", "必须展示样本量，禁止只展示胜率。"),
        N("P4-N05", "计算席位共现特征", "分析席位组合与主题共现后的历史表现。", "seat-core/cooccurrence/", "组合键、支持度、收益分布和置信区间。", "组合顺序归一、稀疏样本和多席位事件处理正确。", "防止组合爆炸与选择偏差。"),
        N("P4-N06", "定义 KOL 内容与 Claim 模型", "区分 Mention、Claim、Evidence、PositionClaim 与事实。", "kol-core/models/、contracts/kol/", "证据等级、方向、时间、证券和不确定性字段。", "模糊表达、转述、回顾性陈述和删除内容测试通过。", "绝不能把 Claim 命名或展示为真实 Trade。"),
        N("P4-N07", "建立 ASR/OCR 内容管线", "从公开视频、音频和截图提取可定位文本。", "intelligence-engine/pipelines/media/", "媒体清单、时间码/坐标证据、质量分和重试。", "低质量、无声、重复帧和超长媒体路径可恢复。", "保留版权、来源和删除策略元数据。"),
        N("P4-N08", "实现 KOL Claim 抽取与审核", "抽取买入、卖出、持仓、关注和否定声明。", "kol-core/extraction/、apps/api/modules/review/", "结构校验、证据定位、置信度和人工复核。", "否定、条件句、引用他人和时间指代样本通过。", "低置信度 Claim 不得进入高等级特征。"),
        N("P4-N09", "计算 KOL 绩效与共识特征", "按公开时点跟踪后续收益、注意力与共识。", "kol-core/performance/、feature-core/kol/", "kol_attention_score、consensus_score、claim_score。", "发布时间、可交易时点、停牌和删除内容口径明确。", "防止用内容发布日期之后的修订文本回填。"),
        N("P4-N10", "构建席位与 KOL 工作台", "展示档案、证据、统计、共现和审核队列。", "apps/web/features/seat-kol/、apps/api/modules/seat-kol/", "shadcn/ui 详情页、时间轴、筛选、样本量和证据视图。", "所有指标显示版本、样本数和数据截止时间。", "UI 不得暗示 KOL Claim 是券商成交事实。"),
    )),
    Phase(5, "Feature Store、量化研究与回测", "统一特征定义，保证点时正确性，并建立符合 A 股与场内基金约束的研究和事件驱动回测能力。", "在线/离线特征一致；数据集可复现；贵金属 NAV/iNAV 不被视为成交价；回测无未来函数并通过会计、订单和回归测试。", (
        N("P5-N01", "定义 FeatureSpec 注册表", "统一特征名称、实体、频率、版本、依赖和有效时点。", "feature-core/registry/、contracts/feature/", "FeatureSpec、注册验证和依赖图。", "重复名称、循环依赖、版本冲突和缺少所有者会失败。", "特征语义变更必须升版本。"),
        N("P5-N02", "实现批量特征计算", "从 ClickHouse/Parquet 生成历史特征。", "feature-core/batch/", "分区计算、检查点、元数据和幂等写入。", "同一输入与版本产生相同输出和校验和。", "避免在计算中隐式读取未来分区。"),
        N("P5-N03", "实现实时特征计算", "消费行情和领域事件生成 Redis 在线特征。", "feature-core/online/", "事件时间窗口、迟到数据、TTL 和降级策略。", "乱序、重复、迟到与重启恢复测试通过。", "处理时间不能替代事件时间。"),
        N("P5-N04", "建立历史特征存储", "在 ClickHouse 保存版本化、点时可查询的特征。", "infra/clickhouse/feature-migrations/、feature-core/storage/", "实体/时间/版本排序与批量查询接口。", "as-of 查询不会返回未来可见值。", "主键与分区需控制高基数。"),
        N("P5-N05", "建立研究数据集导出", "导出带清单和校验和的 Parquet 数据集。", "feature-core/datasets/、research/datasets/", "DatasetManifest、分区、来源和可重建命令。", "任意实验可由清单重建并核对哈希。", "训练与评测分区不得混淆。"),
        N("P5-N06", "建立点时正确性测试库", "自动检测未来函数、修订泄漏和幸存者偏差。", "feature-core/tests/point_in_time/", "合成场景、冻结时钟和反例测试。", "注入未来数据时测试必须失败。", "仅比较时间戳不足以发现修订数据泄漏。"),
        N("P5-N07", "建立因子研究工具", "支持分层、IC、Rank IC、相关性、衰减和换手。", "quant-core/research/factors/", "统一结果对象、图表数据和统计检验。", "常数、缺失、极值和小样本处理有测试。", "报告必须同时展示样本与不确定性。"),
        N("P5-N08", "建立事件研究工具", "研究事件前后窗口的收益与横截面对照。", "quant-core/research/events/", "窗口定义、基准、聚类标准误和输出契约。", "事件重叠、停牌和不可交易样本处理明确。", "避免把事件发生后信息用于事件筛选。"),
        N("P5-N09", "定义策略与信号 SDK", "让策略只读取点时数据并产生 Signal。", "quant-core/strategy/、contracts/strategy/", "Strategy、Context、Signal、参数和生命周期接口。", "策略无法访问未来数据或券商 API。", "SDK 不得暴露 Execution 依赖。"),
        N("P5-N10", "实现事件驱动回测内核", "处理时钟、事件队列、策略、组合和订单意图。", "backtest-core/engine/", "确定性事件循环、随机种子和运行清单。", "同一清单重复运行结果逐项一致。", "避免隐藏的全局状态和真实当前时间。"),
        N("P5-N11", "实现 A 股订单与成交模拟", "模拟 T+1、整手、涨跌停、停牌、费用、滑点和部分成交。", "backtest-core/execution/", "Order、Fill、Fee、Slippage 和成交优先级模型。", "每条交易约束都有独立失败测试和组合测试。", "回测成交价不得使用不可成交价格。"),
        N("P5-N12", "实现组合会计与绩效", "管理现金冻结、持仓成本、公司行动和净值。", "backtest-core/portfolio/、quant-core/metrics/", "复式校验、PnL 分解、回撤、Sharpe/Sortino/Calmar。", "现金+持仓+费用恒等式在每个事件后成立。", "复权研究价不能直接修改真实现金流。"),
        N("P5-N13", "建立回测回归与报告", "保存可复现运行并生成对比报告。", "services/quant-engine/、apps/api/modules/backtest/、apps/web/features/backtest/", "运行清单、基准结果、差异告警和 shadcn/ui 报告页。", "算法或依赖变化造成结果漂移时 CI 能提示。", "不要用自动更新基准掩盖回归。"),
        N("P5-N14", "实现贵金属基金特征与回测夹具", "构建点时正确的溢折价、跟踪、流动性及宏观敏感度特征。", "feature-core/precious-metals/、backtest-core/tests/fixtures/precious-metals/", "版本化 FeatureSpec、基准序列依赖、缺失降级、数据集清单和场内基金回测样例。", "未来 NAV 修订、过期 iNAV、缺少基准和不可成交参考价反例测试通过。", "不得模拟期货保证金、换月或交割，也不得用 NAV/iNAV 作为成交价。"),
    )),
    Phase(6, "智能决策终端与模拟组合", "把市场、贵金属基金、事件、席位、KOL、特征和策略结果整合为可解释终端，并在模拟账户完成完整交易闭环。", "终端可用、贵金属基金数据时效/来源清晰、模拟链路可恢复与对账；所有建议可解释；系统仍不能触达真实账户。", (
        N("P6-N01", "建立终端信息架构", "定义导航、路由、权限和跨模块筛选上下文。", "apps/web/app/、docs/ui/", "站点地图、路由契约、响应式导航和权限矩阵。", "键盘导航、窄屏、无权限和深链接测试通过。", "避免按后端服务名称组织用户导航。"),
        N("P6-N02", "实现市场总览", "展示指数、市场宽度、成交额、情绪和风险状态。", "apps/web/features/dashboard/、apps/api/modules/dashboard/", "可组合卡片、刷新策略和数据时效标识。", "延迟、断流、空数据和历史模式显示明确。", "不要把过期数据展示成实时。"),
        N("P6-N03", "实现涨跌停与连板终端", "展示池、梯队、晋级率和个股证据。", "apps/web/features/limit-pool/", "shadcn/ui 表格、筛选、详情抽屉和导出。", "大数据量虚拟化、排序和无障碍检查通过。", "展示口径必须匹配算法版本。"),
        N("P6-N04", "实现轮动与热点终端", "展示板块生命周期、交叉事件和事件证据。", "apps/web/features/rotation/", "时间轴、对比图、热点详情和历史回放。", "图表数据与表格数据同源且时区一致。", "视觉动画不能掩盖数据断点。"),
        N("P6-N05", "实现股票 Intelligence Ranking", "按特征组合生成可解释排行。", "quant-core/ranking/、apps/web/features/ranking/", "版本化 RankingSpec、分项贡献和筛选。", "相同快照排名稳定；缺失特征处理可解释。", "不得把排行表达为收益保证。"),
        N("P6-N06", "实现个股 Intelligence 详情", "聚合行情、情绪、板块、事件、席位、KOL 和特征。", "apps/web/features/instrument/、apps/api/modules/instrument/", "点时一致的详情查询和证据链接。", "历史时间旅行查询不会混入最新资料。", "避免 N+1 查询和跨库不一致快照。"),
        N("P6-N07", "实现自选与预警规则", "支持用户组合条件、静默期和通知。", "apps/api/modules/watchlist/、notification-worker/、apps/web/features/alerts/", "规则 DSL、去重、节流、历史记录和预览。", "重复事件、时区、禁用规则和失败通知测试通过。", "规则表达式必须受限，不能执行任意代码。"),
        N("P6-N08", "建立模拟账户与组合", "建模 PaperAccount、现金、持仓和净值。", "contracts/portfolio/、quant-core/paper/、apps/api/modules/paper/", "账户生命周期、初始资金和不可变流水。", "任何时点可从流水重建账户状态。", "模拟和未来实盘表必须逻辑隔离。"),
        N("P6-N09", "实现 Signal 到 OrderIntent", "把策略信号转换为目标仓位和订单意图。", "quant-core/portfolio-construction/", "组合约束、舍入、解释和拒绝原因。", "现金不足、整手、集中度和冲突信号测试通过。", "OrderIntent 不能直接发送到券商。"),
        N("P6-N10", "实现模拟成交与订单状态机", "模拟创建、检查、提交、部分成交、成交与撤销。", "quant-core/paper-execution/", "幂等 clientOrderId、状态转换和事件日志。", "重复请求、乱序回报和撤单竞态测试通过。", "非法状态跳转必须拒绝并告警。"),
        N("P6-N11", "实现模拟风控与对账", "在发单前后校验限额并与组合流水对账。", "quant-core/risk/、services/notification-worker/", "风险规则、拒绝事件、日终对账和差异报告。", "越限、断线、重复成交和账面漂移演练通过。", "风控不得只存在于前端。"),
        N("P6-N12", "完成模拟耐久与恢复测试", "连续运行、注入故障并验证恢复。", "tests/endurance/、docs/runbooks/paper-trading.md", "72 小时基线、故障矩阵、恢复手册和结果报告。", "重启、消息重复、数据库短断、缓存丢失后能对账恢复。", "未通过耐久测试不得进入 Phase 7。"),
        N("P6-N13", "构建贵金属基金专题终端", "展示场内黄金及白银相关基金的行情、净值、溢折价、流动性、事件和排行。", "apps/web/features/precious-metals/、apps/api/modules/precious-metals/", "shadcn/ui 产品表、对比视图、事件时间轴、自选预警、缺失和过期状态。", "数据时效、来源、版本、币种和不可用字段清晰；键盘、暗色和响应式检查通过。", "不得暗示 NAV/iNAV 可成交，也不得展示现货/期货交易入口。"),
    )),
    Phase(7, "QMT、独立风控与实盘交易", "在独立部署边界内接入 QMT，并通过多重门禁、模拟优先、对账、风控和小资金灰度安全开放实盘。", "所有法规/券商要求已复核；纸面与真实环境隔离；灾难演练通过；人工批准后才允许小资金实盘。", (
        N("P7-N01", "复核监管与券商准入", "确认届时有效的程序化交易、接口、频率和报告要求。", "docs/compliance/、docs/adr/", "日期化核查记录、券商确认、责任人和禁止项。", "证据来源、核查日期和未决限制完整。", "该节点需人工确认，不能由历史文档代替。"),
        N("P7-N02", "建立独立 Execution Service", "物理隔离执行入口、凭证和部署权限。", "services/execution/、infra/execution/", "独立镜像、网络策略、最小权限和只读健康接口。", "Phase 0–6 服务无法访问其密钥或内部端口。", "控制面身份不能自动拥有交易权限。"),
        N("P7-N03", "定义 BrokerAdapter 与 QMT 契约", "统一账户、持仓、订单、成交和错误语义。", "contracts/broker/、execution/brokers/qmt/", "能力矩阵、错误分类、幂等键和模拟实现。", "契约测试覆盖断线、超时、拒单和重复回报。", "不要泄漏 QMT SDK 类型到领域层。"),
        N("P7-N04", "接入 QMT 模拟环境", "只连接模拟账户并完成查询和订单生命周期。", "execution/brokers/qmt/paper/", "连接、重连、查询、下单、撤单和回报适配。", "配置无法指向真实账户；完整模拟流程测试通过。", "任何实盘凭证都不得用于此节点。"),
        N("P7-N05", "实现持久化订单状态机", "保证订单在进程重启和回报乱序下保持一致。", "execution/orders/、apps/api/migrations/orders/", "Outbox/Inbox、状态转换、审计和 clientOrderId。", "崩溃点注入后无重复下单，非法转换被拒绝。", "数据库提交与外部调用的边界需明确。"),
        N("P7-N06", "实现独立实时风控引擎", "在执行服务外独立计算账户、证券、频率和损失限额。", "services/risk-engine/、contracts/risk/", "规则版本、快照、决策证据和 fail-closed 策略。", "风控不可用时禁止新订单；边界和并发测试通过。", "策略服务不能绕过风控。"),
        N("P7-N07", "实现账户持仓与成交对账", "以券商事实校准内部账本并隔离差异。", "execution/reconciliation/", "启动、周期、日终对账和差异工作流。", "断线、手工交易、部分成交和重复成交演练通过。", "发现差异时必须暂停相关发单。"),
        N("P7-N08", "实现 Kill Switch", "独立禁止新订单、撤销未成交并保留审计。", "services/risk-engine/kill-switch/、apps/web/features/trading-safety/", "多通道触发、权限、确认、状态广播和恢复流程。", "策略服务宕机时仍可触发；恢复需独立审批。", "Kill Switch 不能依赖同一故障域。"),
        N("P7-N09", "建立实盘配置与审批门禁", "用显式环境、双人审批和限额解锁真实账户。", "infra/environments/live/、apps/api/modules/approvals/", "默认关闭、短期授权、审计和自动过期。", "缺少任一配置或审批均安全失败。", "禁止通过配置缺失自动回退实盘。"),
        N("P7-N10", "建立交易可观测性与告警", "监控连接、行情延迟、订单延迟、拒单、PnL 和风控。", "infra/grafana/trading/、notification-worker/trading/", "SLO、仪表盘、分级告警和处置手册。", "告警演练能到达责任人且包含关联 ID。", "指标标签不得暴露账户或密钥。"),
        N("P7-N11", "完成灾难与恢复演练", "演练断网、QMT 重启、数据库切换、时钟漂移和重复消息。", "tests/disaster/、docs/runbooks/live/", "演练脚本、恢复时间、差异报告和改进项。", "每种场景满足 RTO/RPO 和不重复下单要求。", "演练必须在隔离环境且使用模拟账户。"),
        N("P7-N12", "执行小资金实盘灰度", "在人工批准后以极低限额验证真实链路。", "docs/releases/live-pilot/、risk-engine/config/live-pilot/", "限定证券、时段、金额、订单数和退出条件。", "每个交易日人工复盘、券商对账一致、任一异常自动停止。", "该节点不是策略收益验证；不得自动扩大资金。"),
    )),
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_dxa = sum(int(Cm(w).twips) for w in widths_cm)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(Cm(w).twips)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(int(Cm(widths_cm[idx]).twips)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=10.5, bold=False, color=INK, font="Microsoft YaHei", italic=False) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run(run, size=9, color=MUTED)


def add_para(doc, text: str = "", *, size=10.5, bold=False, color=INK, after=6, before=0, align=None, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    if align is not None:
        p.alignment = align
    if text:
        set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, text: str, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    return add_para(doc, text, after=3, style=style)


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size={1: 17, 2: 13.5, 3: 11.5}[level], bold=True, color={1: NAVY, 2: BLUE, 3: TEAL}[level])
    return p


def add_callout(doc, label: str, text: str, fill=LIGHT, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [16.0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(f"{label}  "), bold=True, color=accent)
    set_run(p.add_run(text), color=INK)
    add_para(doc, "", after=3)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float], small=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, value in enumerate(headers):
        set_cell_shading(header.cells[i], NAVY)
        p = header.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(value), size=8.5 if small else 9.5, bold=True, color=WHITE)
    for r_idx, values in enumerate(rows):
        row = table.add_row()
        for c_idx, value in enumerate(values):
            if r_idx % 2:
                set_cell_shading(row.cells[c_idx], PALE)
            p = row.cells[c_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if c_idx == 0 and len(headers) > 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(value), size=8.2 if small else 9.2)
    add_para(doc, "", after=3)
    return table


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for lvl, size, color, before, after in ((1, 17, NAVY, 16, 8), (2, 13.5, BLUE, 12, 6), (3, 11.5, TEAL, 8, 4)):
        s = styles[f"Heading {lvl}"]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        s = styles[name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(10.5)
        s.paragraph_format.space_after = Pt(3)
        s.paragraph_format.line_spacing = 1.15

    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(hp.add_run("昆仑智策  |  KUNLUN ALPHA"), size=8.5, bold=True, color=MUTED)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(fp.add_run("项目总体规划与分阶段实施手册  ·  "), size=8.5, color=MUTED)
    add_field(fp, "PAGE")


def add_cover(doc: Document) -> None:
    add_para(doc, "PROJECT MASTER PLAN", size=10, bold=True, color=GOLD, after=60, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "昆仑智策", size=32, bold=True, color=NAVY, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "KUNLUN ALPHA", size=18, bold=True, color=BLUE, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "A 股智能投研与量化交易平台", size=15, color=INK, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "项目总体规划与分阶段实施手册", size=14, color=TEAL, after=44, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(doc, "核心原则", "观势 · 知势 · 策势。前七阶段只研究、分析、回测与模拟；真实交易仅在 Phase 7 经风控与人工审批后开放。", fill="F2F6F7", accent=TEAL)
    add_para(doc, "适用协作：WorkBuddy 单节点开发 → Codex 节点审核 → 通过后进入下一节点", size=10.5, bold=True, color=NAVY, after=36, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "版本：1.0  |  基线日期：2026-08-11  |  状态：Approved Planning Baseline", size=9.5, color=MUTED, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "技术版本策略：节点启动时解析并锁定最新稳定版；禁止预发布和浮动依赖。", size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    add_heading(doc, "文档控制", 1)
    add_table(doc, ["项目", "内容"], [
        ["品牌", "昆仑智策 · Kunlun Alpha"],
        ["定位", "A 股与场内贵金属基金智能投研、量化研究、回测、模拟组合与受控实盘平台"],
        ["仓库", "kunlun-alpha"],
        ["主要执行者", "WorkBuddy：单节点开发；Codex：需求、架构、代码与测试审核"],
        ["安全边界", "Phase 0–6 禁止真实券商订单；Phase 7 默认关闭并逐级解锁"],
        ["文件规范", "UTF-8 + LF；Windows、macOS、Linux 三端兼容"],
    ], [3.2, 12.8])

    add_heading(doc, "目录", 1)
    toc = [
        "1. 执行摘要", "2. 产品定位与范围", "3. 核心架构与数据流", "4. 技术栈与 Monorepo",
        "5. 统一开发规范", "6. 数据与领域设计", "7. 分阶段路线图", "8. 原子开发节点",
        "9. WorkBuddy 任务模板", "10. Codex 审核清单", "11. 实盘开放门禁", "附录 A：统一配置", "附录 B：术语表",
    ]
    for item in toc:
        add_bullet(doc, item)
    doc.add_page_break()


def add_overview(doc: Document) -> None:
    add_heading(doc, "1. 执行摘要", 1)
    add_para(doc, "昆仑智策不是单纯行情软件，也不是第一天就自动下单的交易机器人。其核心资产是长期可追溯的“事件 → 热点 → 板块 → 股票 → 席位 → KOL Claim → 特征 → 后续收益”历史数据库，以及建立在同一数据口径上的研究、回测、模拟和风险控制能力。")
    add_callout(doc, "交付策略", "每次只向 WorkBuddy 分派一个 P{phase}-N{sequence} 节点。节点完成后提交证据包，由 Codex 审核；未通过不得领取后续依赖节点。", fill="EDF5F5", accent=TEAL)
    add_heading(doc, "2. 产品定位与范围", 1)
    add_heading(doc, "2.1 产品能力", 2)
    for text in ["统一 A 股证券、ETF 与行情数据底座。", "交易所上市黄金及数据源可识别的白银相关 ETF/基金。", "市场情绪、涨跌停、连板梯队和板块轮动。", "新闻、政策、公告、研报和社交内容的事件与热点 Intelligence。", "龙虎榜席位知识库和 KOL 公开 Claim Intelligence。", "版本化 Feature Store、因子研究、事件研究和 A 股/场内基金约束回测。", "智能决策终端、自选预警与模拟组合。", "最后阶段的 QMT 适配、独立风控、订单状态机、对账和小资金灰度。"]:
        add_bullet(doc, text)
    add_heading(doc, "2.2 明确非目标", 2)
    for text in ["第一版不做全市场 Level-2/Tick 长期存储。", "贵金属现货、实物、期货、保证金、换月、交割和仓单业务不在首版范围。", "不做高频交易、超低延迟撮合或自研时序数据库。", "不同时接入多个券商或多个交易市场。", "AI 不直接决定或提交真实买卖订单。", "不以微服务数量、Kubernetes 或 Kafka 作为工程成熟度指标。"]:
        add_bullet(doc, text)
    add_heading(doc, "2.3 不可破坏的原则", 2)
    add_table(doc, ["原则", "要求", "验证方式"], [
        ["安全隔离", "Phase 0–6 无真实下单调用路径", "依赖扫描、网络策略与集成测试"],
        ["数据可追溯", "每个结果可定位来源、版本和处理链", "审计字段与重放测试"],
        ["点时正确", "只使用决策时刻可见的数据", "未来泄漏反例测试"],
        ["契约优先", "生产者与消费者共享版本化契约", "兼容性检查与生成类型"],
        ["小步交付", "单节点可独立测试、提交、回退", "WorkBuddy 证据包与 Codex 门禁"],
    ], [3.0, 7.5, 5.5])

    add_heading(doc, "3. 核心架构与数据流", 1)
    add_para(doc, "外部数据源经适配器进入原始数据区，再经过规范化、校验、去重和事件总线写入不同职责的存储。领域引擎只消费契约化数据；特征、研究和回测共享同一语义。策略输出 Signal 或 OrderIntent，永远不直接调用券商。")
    add_callout(doc, "主数据流", "Data Source → Provider Adapter → Collector → Normalizer → Validator → Deduplicator → Event Bus → Storage → Domain Engines → Feature Store → Research / Backtest / Decision Terminal → Paper Portfolio → Risk Engine → Execution Gateway → QMT", fill="F3F6F8", accent=BLUE)
    add_heading(doc, "3.1 存储职责", 2)
    add_table(doc, ["组件", "主要职责", "禁止事项"], [
        ["PostgreSQL", "用户、配置、任务、审计、订单与组合业务状态", "不承载全市场高频行情"],
        ["ClickHouse", "行情、分钟快照、特征、事件分析与历史查询", "不作为强事务控制面"],
        ["Redis", "在线缓存、实时特征、锁、租约与短期状态", "不作为唯一事实来源"],
        ["MinIO/COS", "不可变原始对象、Parquet 数据集、模型与报告", "不保存无清单的孤立产物"],
    ], [3.1, 8.1, 4.8])
    add_heading(doc, "3.2 服务与领域边界", 2)
    add_table(doc, ["层", "责任"], [
        ["apps/web", "用户终端；只经 API/WS 使用能力"], ["apps/api", "控制面、权限、配置、查询聚合与任务编排"],
        ["services", "可独立运行的 Engine、Worker、Risk 与 Execution"], ["python/packages", "无部署责任的领域库和算法库"],
        ["packages", "UI、契约、客户端、配置和共享类型"], ["infra", "本地/生产基础设施与可观测性"],
    ], [4.0, 12.0])

    add_heading(doc, "4. 技术栈与 Monorepo", 1)
    add_heading(doc, "4.1 技术选型", 2)
    add_table(doc, ["领域", "选型", "约束"], [
        ["Web", "React + TypeScript + Vite", "严格模式；响应式；可访问性"],
        ["组件", "shadcn/ui + Tailwind CSS", "复用原语；业务组件与基础 UI 分离"],
        ["图表", "TradingView Lightweight Charts", "通过内部适配层接入"],
        ["API", "NestJS", "控制面，不执行重型量化计算"],
        ["引擎", "Python + Polars/DuckDB", "类型完整；批流语义一致"],
        ["工程", "Turborepo + pnpm + uv", "锁文件提交；版本精确"],
        ["可观测", "OpenTelemetry + Prometheus + Grafana + Loki", "统一 request/trace ID"],
    ], [3.0, 6.0, 7.0])
    add_heading(doc, "4.2 目标目录", 2)
    structure = """kunlun-alpha/
├─ apps/                 # web、api
├─ services/             # market/intelligence/quant/data/notification/risk/execution
├─ python/packages/      # common、contracts 与各领域 core
├─ packages/             # ui、contracts、shared-types、api-client、config
├─ infra/                # docker、db、object storage、observability、deployment
├─ research/             # notebooks、experiments、evaluations、reports
├─ data/                 # fixtures 与脱敏 samples
├─ docs/                 # architecture、ADR、runbook、algorithm、compliance
└─ scripts/              # 跨平台开发、CI 和运维入口"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    for line in structure.splitlines():
        r = p.add_run(line + "\n")
        set_run(r, size=8.8, color=INK, font="Cascadia Mono")


def add_standards(doc: Document) -> None:
    add_heading(doc, "5. 统一开发规范", 1)
    sections = {
        "5.1 版本与依赖": ["节点启动时解析最新稳定版，随后写入版本基线并锁定。", "禁止 alpha、beta、RC、nightly 与 Docker latest。", "JavaScript 使用精确版本并提交 pnpm-lock.yaml；Python 提交 uv.lock。", "依赖升级作为独立节点，必须有变更、兼容性、测试和回滚说明。"],
        "5.2 LF 与三端兼容": ["除显式 Windows 批处理外，全部文本使用 UTF-8 + LF。", "使用跨平台路径 API，不硬编码盘符、斜杠、主目录或 /tmp。", "脚本优先采用跨平台 Node/Python；三端入口命令保持一致。", "CI 至少 Linux 全量，Windows/macOS 验证 bootstrap、lint、typecheck 和核心测试。"],
        "5.3 TypeScript": ["tsconfig 开启 strict、noUncheckedIndexedAccess 等安全选项。", "禁止无理由 any；外部输入先以 unknown 接收并验证。", "API 请求封装在生成客户端或 feature service，不在页面拼接 URL。", "模块导出通过显式 public API，禁止深层路径穿透。"],
        "5.4 Python": ["公共函数、类和领域边界使用完整类型标注并通过 Pyright。", "I/O 与纯计算分离；算法优先纯函数和不可变输入。", "金额与价格使用 Decimal/整数最小单位；时间使用 timezone-aware datetime。", "禁止把 Notebook 代码直接复制为生产服务。"],
        "5.5 测试": ["默认 TDD：失败测试 → 最小实现 → 通过 → 重构。", "每节点覆盖正常、边界、失败三类路径。", "外部 Provider、对象存储、LLM 和券商提供契约化模拟实现。", "测试禁止真实网络、真实账户和不可控当前时间。"],
        "5.6 Git 与评审": ["Conventional Commits；提交聚焦并在正文引用节点 ID。", "禁止无关格式化、依赖升级或大范围重构混入业务节点。", "提交前运行格式、静态分析、窄测试和受影响包测试。", "WorkBuddy 提交证据包；Codex 审核通过后才进入依赖节点。"],
        "5.7 安全与隐私": ["密钥只进入 Secret 管理；示例配置使用不可用占位值。", "日志默认脱敏账户、Token、Cookie、手机号与原始敏感内容。", "权限最小化；模拟与实盘凭证、数据库和网络边界分离。", "依赖、镜像和许可证扫描进入 CI。"],
        "5.8 可观测性": ["日志结构化并包含 service、version、request_id、trace_id。", "指标控制基数；证券代码通常不作为长期 Prometheus 标签。", "任务、数据延迟、队列、错误率与外部 Provider 限流均有告警。", "Runbook 与告警规则同时交付。"],
        "5.9 shadcn/ui": ["基础组件只从 components/ui 复用或按 shadcn 方式纳入代码库。", "业务组件不修改基础原语的语义和可访问性。", "所有页面提供 loading、empty、error、forbidden 状态。", "暗色、键盘、焦点、窄屏和高密度表格纳入视觉验收。"],
    }
    for title, bullets in sections.items():
        add_heading(doc, title, 2)
        for b in bullets:
            add_bullet(doc, b)

    add_heading(doc, "5.10 Definition of Done", 2)
    add_table(doc, ["门禁", "通过条件"], [
        ["范围", "仅实现当前节点，非目标未被提前带入"],
        ["代码", "边界清晰、类型完整、无秘密、无平台硬编码"],
        ["测试", "失败证据、通过证据、边界/失败路径完整"],
        ["质量", "format、lint、typecheck、test、build 全部通过"],
        ["文档", "契约、ADR、配置样例、Runbook 与变更同步"],
        ["审核", "WorkBuddy 证据包完整，Codex 无未解决阻断问题"],
    ], [3.2, 12.8])


def add_domain_design(doc: Document) -> None:
    add_heading(doc, "6. 数据与领域设计", 1)
    rows = [
        ["Instrument", "symbol、exchange、board、list/delist、ST、limit rule", "PostgreSQL + 合约"],
        ["MarketEmotionSnapshot", "涨跌停、炸板、连板、晋级率、溢价、score_version", "ClickHouse"],
        ["SectorSnapshot", "涨速、成交额、宽度、龙头、rotation/hot/emotion score", "ClickHouse"],
        ["RotationEvent", "from、to、event_time、confidence、evidence window", "ClickHouse"],
        ["MarketEvent", "type、sentiment、impact、topics、stocks、evidence、model_version", "PostgreSQL/ClickHouse"],
        ["HotTopic", "窗口聚合、HotScore、生命周期、证据", "ClickHouse"],
        ["SeatProfile", "stable seat_id、alias、trade、performance、preference", "PostgreSQL/ClickHouse"],
        ["KolClaim", "claim type、direction、time、evidence、confidence", "PostgreSQL/对象存储"],
        ["PreciousMetalsFund", "fundAssetClass、underlying commodity、NAV/iNAV、premium/discount、benchmark", "PostgreSQL/ClickHouse"],
        ["FeatureValue", "entity、event_time、available_time、name、version、value", "Redis/ClickHouse/Parquet"],
        ["OrderIntent", "strategy、account、symbol、side、qty/target、reason", "业务数据库"],
    ]
    add_table(doc, ["核心对象", "关键语义", "主要落点"], rows, [3.5, 8.5, 4.0], small=True)
    add_callout(doc, "时间语义", "必须区分 event_time、publish_time、ingest_time、available_time 和 processing_time。回测与特征 as-of 查询只允许使用 available_time 不晚于决策时刻的数据。", fill="FFF7EA", accent=GOLD)
    add_callout(doc, "版本语义", "emotion_score_v1、HotScore V1、FeatureSpec、Prompt、Schema、DatasetManifest 和回测运行清单均不可变。语义改变必须新增版本。", fill="EEF4F8", accent=BLUE)
    add_callout(doc, "贵金属范围", "首版只支持交易所上市的黄金及数据源可识别的白银相关 ETF/基金。NAV/iNAV 是点时研究参考值，不是成交价；现货、期货、保证金、换月、交割和实物业务必须另行设计。", fill="FFF7EA", accent=GOLD)


def add_phase_roadmap(doc: Document) -> None:
    add_heading(doc, "7. 分阶段路线图", 1)
    rows = []
    for phase in PHASES:
        rows.append([f"Phase {phase.number}", phase.name, str(len(phase.nodes)), phase.exit_gate])
    add_table(doc, ["阶段", "核心目标", "节点数", "退出门槛"], rows, [2.0, 4.5, 1.8, 7.7], small=True)
    total = sum(len(p.nodes) for p in PHASES)
    add_callout(doc, "计划规模", f"共 {total} 个原子开发节点。节点通常为 0.5–2 个开发日；实际排期应在领取节点时依据代码现状重新估算，不得把整阶段一次性交给 WorkBuddy。", fill="F3F6F8", accent=NAVY)


def add_node_card(doc: Document, node: Node, predecessor: str, phase: Phase) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(node.node_id), size=10, bold=True, color=WHITE)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), TEAL if phase.number < 7 else RED)
    p_pr.append(shd)
    set_run(p.add_run(f"  {node.title}"), size=11, bold=True, color=WHITE)

    rows = [
        ["目标", node.objective],
        ["依赖", predecessor],
        ["文件区域", node.files],
        ["主要交付", node.deliverable],
        ["实施循环", "编写失败测试/验证夹具 → 运行并确认失败 → 最小实现 → 窄测试通过 → 受影响包全测 → 更新文档与提交。"],
        ["验收", node.acceptance],
        ["审核重点", node.risk],
        ["证据包", f"{node.node_id} 变更摘要、文件清单、测试命令与输出、静态检查/构建结果、截图或样例载荷、风险和未决事项。"],
        ["提交建议", f"feat({node.node_id.lower()}): {node.title}"],
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [2.6, 13.4])
    for idx, (label, value) in enumerate(rows):
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT)
        if idx % 2:
            set_cell_shading(row.cells[1], PALE)
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_after = Pt(0)
        set_run(p0.add_run(label), size=8.5, bold=True, color=NAVY)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing = 1.08
        set_run(p1.add_run(value), size=8.5)
    add_para(doc, "", after=2)


def add_nodes(doc: Document) -> None:
    add_heading(doc, "8. 原子开发节点", 1)
    add_para(doc, "本章是 WorkBuddy 的权威任务目录。默认采用阶段内顺序依赖；若节点明确只依赖更早公共契约，可在 Codex 确认后并行。任何并行不得共享未稳定接口。")
    for phase in PHASES:
        doc.add_page_break()
        add_heading(doc, f"8.{phase.number + 1} Phase {phase.number} — {phase.name}", 1)
        add_para(doc, phase.goal, bold=True, color=NAVY)
        add_callout(doc, "阶段退出门槛", phase.exit_gate, fill="F4F7F8", accent=TEAL if phase.number < 7 else RED)
        prev = "已批准的项目设计稿；无代码前置依赖" if phase.number == 0 else f"Phase {phase.number - 1} 退出门槛已通过 Codex 审核"
        for node in phase.nodes:
            add_node_card(doc, node, prev, phase)
            prev = f"{node.node_id} 已通过 Codex 审核"


def add_workflow(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "9. WorkBuddy 单节点任务模板", 1)
    template_rows = [
        ["任务", "只执行节点：<NODE_ID> <TITLE>"],
        ["上下文", "阅读本手册的全局约束、所属 Phase、当前节点与直接依赖节点；不要扩展范围。"],
        ["实施", "先写失败测试并记录失败，再完成最小实现；所有外部依赖使用模拟或测试环境。"],
        ["禁止", "不得领取下一节点；不得顺带升级依赖、重构无关模块或连接真实券商。"],
        ["验证", "依次运行 format、lint、typecheck、narrow test、affected test、build。"],
        ["交付", "变更摘要、文件清单、接口说明、测试输出、截图/样例、风险、未决事项和提交哈希。"],
    ]
    add_table(doc, ["字段", "交给 WorkBuddy 的固定要求"], template_rows, [3.2, 12.8])
    add_heading(doc, "9.1 标准任务提示词", 2)
    prompt = """请仅实施《昆仑智策项目总体规划与分阶段实施手册》中的 <NODE_ID>：<TITLE>。
严格遵守全局开发规范、LF 与三端兼容要求，以及该节点的依赖、文件区域、验收和审核重点。
先编写失败测试并确认失败，再完成最小实现。不得实现后续节点，不得进行无关重构或依赖升级。
完成后提供：变更摘要、文件清单、关键设计决定、全部验证命令与结果、截图或样例载荷、风险、未决事项、提交哈希。然后停止，等待 Codex 审核。"""
    p = doc.add_paragraph()
    set_cell = None
    for line in prompt.splitlines():
        set_run(p.add_run(line + "\n"), size=9, color=INK, font="Cascadia Mono")

    add_heading(doc, "10. Codex 审核清单", 1)
    checks = [
        ["需求", "是否完整满足当前节点，且没有提前实现后续范围？"],
        ["边界", "依赖方向、契约、存储职责和服务责任是否正确？"],
        ["类型", "TypeScript/Python 类型是否明确，外部输入是否验证？"],
        ["数据", "时间、精度、版本、幂等、审计和点时语义是否正确？"],
        ["测试", "是否先失败后通过，覆盖正常、边界、失败及关键回归？"],
        ["安全", "是否泄密、越权、暴露实盘路径或出现 fail-open？"],
        ["跨平台", "是否保持 LF，路径、脚本和命令是否适配三端？"],
        ["前端", "是否复用 shadcn/ui，覆盖四态、暗色、键盘和响应式？"],
        ["运维", "日志、指标、告警、Runbook 和回滚是否与风险匹配？"],
        ["提交", "变更是否聚焦，文档与证据包是否完整？"],
    ]
    add_table(doc, ["检查域", "审核问题"], checks, [3.2, 12.8])
    add_heading(doc, "10.1 审核结论", 2)
    for text in ["通过：无阻断问题，可进入后续依赖节点。", "有条件通过：只有不影响正确性和安全性的明确跟踪项。", "退回：需求、架构、数据正确性、测试或安全存在阻断问题。", "停止：发现真实交易绕过、密钥泄露、不可逆数据破坏或无法解释的对账差异。"]:
        add_bullet(doc, text)


def add_live_gate_and_appendix(doc: Document) -> None:
    add_heading(doc, "11. 实盘开放门禁", 1)
    gates = [
        ["法规与券商", "核实届时有效规则、程序化报告、QMT 权限、接口与频率限制", "人工签字"],
        ["环境隔离", "模拟/实盘账户、凭证、数据库、网络、日志和审批完全隔离", "安全审计"],
        ["订单安全", "幂等、状态机、重复回报、撤单竞态、崩溃恢复测试通过", "测试报告"],
        ["独立风控", "风控不可用时 fail-closed；策略无法绕过", "故障注入"],
        ["对账", "启动、周期、日终对账与差异暂停机制通过", "演练证据"],
        ["Kill Switch", "独立故障域、多通道触发、恢复需审批", "实操演练"],
        ["可观测性", "连接、延迟、拒单、PnL、风控与告警链路完整", "值守确认"],
        ["灰度", "限定证券、时段、金额、订单数和停止条件", "逐日批准"],
    ]
    add_table(doc, ["门禁", "要求", "证据"], gates, [3.0, 9.5, 3.5], small=True)
    add_callout(doc, "默认结论", "只要任一门禁证据缺失，系统结论必须是“禁止实盘”。不得以配置缺失、负责人不在线或市场机会紧迫为理由绕过。", fill="FBEFEF", accent=RED)

    add_heading(doc, "附录 A：统一配置", 1)
    add_heading(doc, "A.1 .editorconfig", 2)
    editorconfig = """root = true

[*]
charset = utf-8
end_of_line = lf
insert_final_newline = true
indent_style = space
indent_size = 2
trim_trailing_whitespace = true

[*.py]
indent_size = 4

[*.{md,mdx}]
trim_trailing_whitespace = false

[*.{yaml,yml}]
indent_size = 2

[Makefile]
indent_style = tab

[*.{bat,cmd}]
end_of_line = crlf"""
    p = doc.add_paragraph()
    for line in editorconfig.splitlines():
        set_run(p.add_run(line + "\n"), size=8.5, font="Cascadia Mono", color=INK)
    add_heading(doc, "A.2 .gitattributes", 2)
    attrs = """* text=auto eol=lf
*.bat text eol=crlf
*.cmd text eol=crlf
*.png binary
*.jpg binary
*.jpeg binary
*.gif binary
*.pdf binary
*.docx binary
*.parquet binary"""
    p = doc.add_paragraph()
    for line in attrs.splitlines():
        set_run(p.add_run(line + "\n"), size=8.5, font="Cascadia Mono", color=INK)

    add_heading(doc, "附录 B：术语表", 1)
    terms = [
        ["Provider", "外部数据或服务的适配接口，领域逻辑不依赖具体供应商。"],
        ["Claim", "KOL 公开表达的交易或持仓声明，不等同券商成交事实。"],
        ["Point-in-time", "只使用目标时刻已经可见的数据进行研究或决策。"],
        ["FeatureSpec", "特征名称、实体、频率、版本、依赖与有效时间的定义。"],
        ["OrderIntent", "经组合构建得到、尚未进入券商的订单意图。"],
        ["Fail-closed", "依赖或风控异常时默认拒绝新订单，而非继续交易。"],
        ["Kill Switch", "独立禁止新订单并可撤销未成交订单的安全机制。"],
        ["Evidence Packet", "WorkBuddy 节点交付时提供的代码、测试、截图、风险和提交证据。"],
    ]
    add_table(doc, ["术语", "定义"], terms, [4.0, 12.0])


def build() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_overview(doc)
    add_standards(doc)
    add_domain_design(doc)
    add_phase_roadmap(doc)
    add_nodes(doc)
    add_workflow(doc)
    add_live_gate_and_appendix(doc)
    props = doc.core_properties
    props.title = "昆仑智策项目总体规划与分阶段实施手册"
    props.subject = "Kunlun Alpha A-share Intelligence & Quant Platform"
    props.author = "Kunlun Alpha Project"
    props.keywords = "A股, 量化, Intelligence, WorkBuddy, Codex, QMT"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)
    print(f"phases={len(PHASES)} nodes={sum(len(p.nodes) for p in PHASES)}")


if __name__ == "__main__":
    build()
